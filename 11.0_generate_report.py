import os
import json
import re
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants

# --- Load Environment Variables ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

zai_key = os.getenv("ZAI_API_KEY")
if not zai_key:
    raise ValueError("ZAI_API_KEY not found. Please check your .env file.")

client = OpenAI(
    api_key=zai_key,
    base_url="https://api.z.ai/api/paas/v4/"
)

def make_heading_collapsible(heading):
    """Injects Word XML tags to make a heading load in a collapsed state by default."""
    pPr = heading._element.get_or_add_pPr()
    collapsed = OxmlElement('w:collapsed')
    collapsed.set(qn('w:val'), '1')
    pPr.append(collapsed)

def add_hyperlink(paragraph, text, url):
    """Adds a clickable hyperlink to a paragraph in python-docx."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000EE')
    rPr.append(c)
    
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    
    text_elem = docx.oxml.shared.OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def load_json_safe(filepath):
    """Helper to load JSON files without crashing if one is missing or malformed."""
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except json.JSONDecodeError as e:
            print(f"  [!] WARNING: {os.path.basename(filepath)} contains invalid JSON and was skipped. (Error: {e})")
            return {}
    return {}

def strip_hallucinated_citations(text):
    """Removes fake LLM citations like [1.2.5] or [3] from the text."""
    if not text:
        return ""
    return re.sub(r'\[[\d\.\s,]+\]', '', text).strip()

def sanitize_to_diplomatic_tone(text):
    """
    Automated Sanitizer Layer: Automatically maps critical or aggressive phrases 
    into highly polished, constructive executive language.
    """
    if not text:
        return ""
    
    replacements = {
        "highly restricted or non-existent digital footprint": "highly curated and selective digital footprint",
        "lacking even a basic corporate LinkedIn presence": "maintaining a stealth profile on mainstream professional networks",
        "This complete absence across major platforms indicates": "This private, relationship-driven posture suggests",
        "lacking even": "currently opting away from",
        "complete absence": "highly specialized focus",
        "non-existent": "nascent",
        "deficient": "unoptimized"
    }
    
    for harsh_word, gentle_word in replacements.items():
        text = re.sub(re.escape(harsh_word), gentle_word, text, flags=re.IGNORECASE)
    return text

def synthesize_ai_strategy(company_name, llm_insights, industry_ai):
    """Uses GLM-5.2 to extract score and prioritized recommendations."""
    system_prompt = """
    You are an expert AI implementation strategist and elite executive consultant. 
    You are provided with data scraped from a target company's website and real-world AI use cases for their industry.
    
    CRITICAL TONE INSTRUCTION:
    Your tone must be highly professional, diplomatic, and constructive. 
    DO NOT use aggressive or negative words like 'deficient', 'stagnant', 'lacking', or 'failing'. 
    Frame all technical gaps as 'open strategic horizons', 'unrealized optimization opportunities', or 'nascent integration phases'.
    
    You must output ONLY valid JSON format with the following structure:
    {
        "score": integer (0 to 100 representing their current AI adoption maturity),
        "score_justification": [array of 3-4 string bullet points backing up the score],
        "recommendations": [array of highly actionable string bullet points on how AI can help them. SEQUENCED IN ORDER of strategic importance, from highest ROI/urgency to lowest.]
    }
    """
    
    user_prompt = f"Company: {company_name}\n\nCompany Web Insights:\n{json.dumps(llm_insights)}\n\nIndustry AI Context:\n{json.dumps(industry_ai)}"

    print(f"\n[+] Asking GLM-5.2 to synthesize AI Strategy for {company_name}...")
    try:
        response = client.chat.completions.create(
            model="glm-5.2",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3 
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"  [!] Error calling Z.AI API: {e}")
        return {"score": 0, "score_justification": ["Error generating score"], "recommendations": ["Error generating recommendations"]}

def build_word_document(folder_path, company_name):
    llm_insights = load_json_safe(os.path.join(folder_path, "llm_insights.json"))
    industry_ai = load_json_safe(os.path.join(folder_path, "industry_ai.json"))
    peer_data = load_json_safe(os.path.join(folder_path, "peer_ai_analysis.json"))
    news_data = load_json_safe(os.path.join(folder_path, "news.json"))
    leadership_data = load_json_safe(os.path.join(folder_path, "leadership.json"))
    social_data = load_json_safe(os.path.join(folder_path, "social_footprint.json"))
    metadata = load_json_safe(os.path.join(folder_path, "domain_metadata.json"))
    hiring_data = load_json_safe(os.path.join(folder_path, "hiring_signals.json"))

    ai_strategy = synthesize_ai_strategy(company_name, llm_insights, industry_ai)

    doc = docx.Document()
    
    # --- Title & Date ---
    title = doc.add_heading(company_name.upper(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.italic = True
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1: AI Readiness Score ---
    doc.add_heading("1. AI Readiness Score", level=1)
        
    p_score = doc.add_paragraph()
    runner = p_score.add_run(f"Score: {ai_strategy.get('score', 'N/A')}/100")
    runner.bold = True
    runner.font.size = Pt(14)
    runner.font.color.rgb = RGBColor(0, 102, 204) 
    
    for bullet in ai_strategy.get('score_justification', []):
        doc.add_paragraph(sanitize_to_diplomatic_tone(strip_hallucinated_citations(bullet)), style='List Bullet')

    if metadata:
        infra = metadata.get("infrastructure", {})
        has_txt = infra.get("has_llms_txt", False)
        has_full = infra.get("has_llms_full_txt", False)
        
        infra_bullet = doc.add_paragraph(style='List Bullet')
        infra_bullet.add_run("AI Discovery Infrastructure: ").bold = True
        if has_txt or has_full:
            infra_bullet.add_run("Optimized. The domain contains modern llms.txt standard files and is positioned to leverage rapid indexing protocols like IndexNow. This makes the company's data highly legible and instantly available to frontier AI models, scrapers, and agents during search grounding tasks.")
        else:
            infra_bullet.add_run("Not Yet Implemented. The domain currently maps to standard consumer traffic and can be further optimized by introducing machine-readable agent endpoints (llms.txt / llms-full.txt). Integrating these alongside immediate ping protocols like IndexNow ensures real-time Legibility across frontier model search arrays.")

    # --- Section 2: How AI Can Help ---
    doc.add_heading("2. Strategic AI Opportunities", level=1)
    for rec in ai_strategy.get('recommendations', []):
        doc.add_paragraph(sanitize_to_diplomatic_tone(strip_hallucinated_citations(rec)), style='List Bullet')

    # --- Section 3: AI Human Capital Investment ---
    doc.add_heading("3. AI Human Capital Investment", level=1)
    
    if hiring_data:
        ai_keywords = r'(artificial intelligence|ai|machine learning|ml|llm|data scientist|deep learning)'
        ai_roles_found = []
        
        for role in hiring_data.get("open_roles", []):
            title_match = re.search(ai_keywords, role.get("job_title", "").lower())
            focus_match = re.search(ai_keywords, role.get("focus_area", "").lower())
            if title_match or focus_match:
                ai_roles_found.append(role)
                
        p_hiring = doc.add_paragraph()
        
        if ai_roles_found:
            p_hiring.add_run("Active AI Talent Acquisition: ").bold = True
            p_hiring.add_run("Current hiring footprints confirm strategic human capital allocation toward core AI integration, notably targeting the following technical benchmarks:")
            
            for role in ai_roles_found:
                p_role = doc.add_paragraph(style='List Bullet')
                add_hyperlink(p_role, role.get("job_title", "Open AI Role"), role.get("url", "#"))
                focus_area = role.get("focus_area")
                if focus_area:
                    p_role.add_run(f" – {focus_area}")
                    
        else:
            if hiring_data.get("is_hiring_tech", False):
                p_hiring.add_run("Hardware & Core Infrastructure Prioritized: ").bold = True
                p_hiring.add_run("Open requisitions reveal robust, active investment heavily prioritized toward technical hardware development and foundational engineering. While these positions form an exceptional baseline, explicit public allocations for dedicated software machine learning or data science roles have not yet been introduced to the talent board.")
            else:
                p_hiring.add_run("Curated Internal Operations: ").bold = True
                p_hiring.add_run("Public-facing talent acquisition does not feature active engineering or specialized artificial intelligence vacancies at this moment, highlighting an open operational horizon to introduce these capabilities as your data roadmap expands.")
    else:
        doc.add_paragraph("No public hiring data or requisitions could be evaluated for this company.")

    # --- Section 4: Industry AI Context ---
    doc.add_heading("4. Industry AI Context", level=1)
        
    if industry_ai:
        doc.add_heading("Current Industry Sector AI Findings", level=2)
        doc.add_paragraph(sanitize_to_diplomatic_tone(strip_hallucinated_citations(industry_ai.get("ai_paragraph", ""))))
        for uc in industry_ai.get("use_cases", []):
            doc.add_paragraph(f"{strip_hallucinated_citations(uc.get('use_case'))}: {sanitize_to_diplomatic_tone(strip_hallucinated_citations(uc.get('details')))}", style='List Bullet')

    # --- Section 5: Peer AI Benchmarking ---
    doc.add_heading("5. Peer AI Benchmarking", level=1)
    if peer_data:
        for peer in peer_data:
            summary = peer.get("ai_initiatives_summary", "")
            if "Could not evaluate AI posture" in summary:
                continue
                
            p_peer = doc.add_heading(level=2)
            add_hyperlink(p_peer, peer.get("company_name", "Peer"), peer.get("homepage_url", ""))
            doc.add_paragraph(sanitize_to_diplomatic_tone(strip_hallucinated_citations(summary)))

    # --- Section 6: Supplemental Intelligence (COLLAPSIBLE) ---
    sup_heading = doc.add_heading("6. Supplemental Intelligence", level=1)
    make_heading_collapsible(sup_heading) 

    # Leadership
    if leadership_data:
        doc.add_heading("Executive Leadership", level=2)
        for exec_member in leadership_data.get("executives", []):
            doc.add_paragraph(f"{exec_member.get('name')} - {exec_member.get('title')} (Tenure: {exec_member.get('tenure')})", style='List Bullet')

    # Recent News
    if news_data:
        doc.add_heading("Recent News & Press", level=2)
        for news in news_data.get("news_items", []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{news.get('date')} | ")
            add_hyperlink(p, news.get('headline', 'News Link'), news.get('url', '#'))

    # Social Footprint
    if social_data:
        doc.add_heading("Digital & Social Footprint", level=2)
        doc.add_paragraph(sanitize_to_diplomatic_tone(social_data.get("digital_presence_summary", "")))
        for channel in social_data.get("channels", []):
            p = doc.add_paragraph(style='List Bullet')
            platform_name = channel.get('platform', 'Social Profile')
            add_hyperlink(p, platform_name, channel.get('profile_url', '#'))
            
            # Clean up the follower count display
            followers = str(channel.get('follower_count', '')).strip()
            if followers.lower() in ['n/a', 'not found', 'unknown', 'none', '']:
                follower_str = ""
            else:
                follower_str = f": {followers} Followers" if platform_name.lower() != 'youtube' else f": {followers} Subscribers"

            # Append the correct metrics based on platform
            if platform_name.lower() == 'youtube':
                video_count = channel.get('video_count', 'N/A')
                p.add_run(f"{follower_str} | Videos: {video_count}")
            else:
                if follower_str:
                    p.add_run(follower_str)

    # Save the file
    output_path = os.path.join(folder_path, f'{company_name.replace(" ", "_")}_Executive_Report.docx')
    doc.save(output_path)
    print(f"\n[!] Success! Final Report compiled and saved to: {output_path}")

if __name__ == "__main__":
    if os.path.exists("config.json"):
        with open("config.json", "r", encoding="utf-8") as f:
            config = json.load(f)
        
        target_folder = config.get("domain_folder")
        company_name = config.get("proper_company_name", target_folder.split('.')[0].capitalize())
        
        print(f"[+] Automating final document generation for: {company_name}")
        build_word_document(target_folder, company_name)
    else:
        print("Error: config.json not found. Run 00_run_pipeline.py first.")